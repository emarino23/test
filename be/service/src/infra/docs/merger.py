import json
import os
from docx import Document
import requests
from naver_core import *
any_module(__file__, 3)
from docxcompose.composer import Composer
from docx.shared import Inches
from docx.shared import Pt
from datetime import datetime
from docx2pdf import convert
from datetime import timedelta
import yaml
from src.infra.web.app.routes import app
import shutil
from naver_net import NaverNet
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.business.Tour.FindCruise import FSFindCruise



class DocGenerator(object):
    def __init__(self, data):
        try:
            self.data = data
            self.customer_data = self.data.get("customer")
            self.tour_data = self.data.get("tour")
            self.logistic_data = self.data.get("logistic")
            self.header = {**self.customer_data, **
                           self.tour_data, **self.logistic_data}
            self.file_dir = os.path.dirname(__file__)
            self.root_dir = os.path.dirname(self.file_dir)
            self.docdir = os.path.join(self.root_dir, ("web\\static\\docs"))
            self.tempdir = os.path.join(self.file_dir, (f"tmp\\"))
            self.weekdays = ["Monday", "Tuesday",
                             "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
            self.email = os.environ.get('EMAIL')
            self.travel_code = self.header.get("travel_code")
            self.destinations = self.data.get("destinations")
            self.until_date = datetime.now().date().replace(month=12, day=31)
            self.passengers = self.header.get("passengers")
            self.arrival_date = self.header.get("arrival_date")
            self.departure_date = self.header.get("departure_date")
            self.names = str(self.header.get("names")).upper()
            self.last_names = str(self.header.get("last_names")).upper()
            self.arrival = self.check_date(
                self.arrival_date, "%Y-%m-%d %H:%M:%S")
            self.departure = self.check_date(
                self.departure_date, "%Y-%m-%d %H:%M:%S")
            self.net = NaverNet(app)
            self.docs = set()
            self.until = self.check_date(
                self.until_date, "%Y-%m-%d %H:%M:%S.%f")
            self.days = (self.departure - self.arrival).days
            self.nights = self.days - 1
            self.customer = f"{self.names} {self.last_names}"
            self.composed = os.path.join(
                self.docdir, (f"{self.travel_code}.docx"))
            self.pdf = os.path.join(self.docdir, (f"{self.travel_code}.pdf"))
            self.delete_doc(self.composed)
            self.delete_doc(self.pdf)
        except Exception as e:
            print(e)
            raise e

    def gen_tour_doc(self):
        try:
            self.gen_cover()
            self.gen_destinations()
            self.gen_doc()
            self.send_email_doc()
            self.del_tmp_folder()
        except Exception as e:
            print(e)
            raise e

    def gen_cover(self):
        try:
            doc_template = os.path.join(self.file_dir, (f"cover.docx"))
            doc_output = os.path.join(
                self.tempdir, (f"cover-{self.travel_code}.docx"))
            document = Document(doc_template)
            self.replace_word(document, "NIGHTS", f"{self.nights}")
            self.replace_word(document, "DAYS", f"{self.days}")
            self.replace_word(document, "CUSTOMER",
                              f"{str(self.customer).upper()}")
            self.replace_word(document, "LEGAL_NAME",
                              f"{str(self.customer).upper()}")
            self.replace_word(document, "PAX", f"{self.passengers}")
            self.replace_word(document, "VALID",
                              f'{self.until.strftime("%Y-%m-%d")}')
            document.save(doc_output)
            self.docs.add(doc_output)
        except Exception as e:
            print(e)
            raise e

    def gen_destinations(self):
        try:
            for dest_id, dest_name in enumerate(self.destinations):
                destination = self.destinations[dest_name]
                dest_name = destination["destination"]
                doc = Document()
                heading = doc.add_heading(
                    f"{str(dest_name).replace('_',' ').capitalize()}", 0)
                run = heading.add_run()
                days = destination.get("daysData")
                self.gen_days(doc, dest_id, dest_name, days)
        except Exception as e:
            print(e)
            raise e

    def gen_cruise(self, doc, dest_name):
        try:
            logistic = self.logistic_data
            if "cruise" in dest_name:
                payload = {
                    "data": {
                        "cruise_name": logistic['cruiseName']
                    }
                }
                cruiseData = FSFindCruise(payload)['data'][0]
                arrival_port = cruiseData['arrival_port']
                image_url = cruiseData['image']

                heading = doc.add_heading(
                    f"{str(logistic['cruiseName'])}  ", 1)
                run = heading.add_run()
                heading = doc.add_heading(
                    f" {arrival_port} ", 1)
                run = heading.add_run()
                image = self.download_image(image_url)
                doc.add_picture(image, width=Inches(4),
                                height=Inches(4))
                paragraph = doc.add_paragraph(
                    f"Itinerary:  {str(cruiseData['cruise_itinerary']).replace('[','').replace(']','')}", "List Paragraph")
                run = paragraph.add_run()
        except Exception as e:
            print(e)
            raise e

    def download_image(self, url):
        try:
            b_data = requests.get(url).content
            file_path = os.path.join(self.tempdir, "criuse.png")
            output = open(file_path, "wb")
            output.write(b_data)
            output.close()
            return file_path
        except Exception as e:
            print(e)
            raise e

    def gen_days(self, doc, dest_id, dest_name, days):
        try:
            for day_id, day_name in enumerate(days):
                weekday, experiences, date_string, = self.prepare_day(
                    days, day_name)
                args = [doc, days, dest_id,  day_id,
                        day_name, date_string, weekday]
                self.gen_meals(*args)
                if "cruise" not in dest_name:
                    self.gen_experiences(doc, dest_name, experiences)
            output = os.path.join(
                self.tempdir, (f"{dest_name}-{self.travel_code}-{dest_id}.docx"))
            self.gen_cruise(doc, dest_name)
            doc.save(output)
            self.docs.add(output)
        except Exception as e:
            print(e)
            raise e

    def prepare_day(self, days, day_name):
        try:
            day = days[day_name]
            experiences = day.get("experiences")
            self.meals = day.get("meals")
            self.type = "Tour"
            date = (self.arrival + timedelta(days=int(day_name))
                    )
            date_string = date.strftime("%Y-%m-%d")
            weekday = self.weekdays[date.weekday()]
            return [weekday, experiences, date_string]
        except Exception as e:
            print(e)
            raise e

    def gen_meals(self, *args):
        try:
            doc, days, dest_id,  day_id, day_name, date_string, weekday, = args
            if (day_id) == 0 and dest_id == 0:
                self.type = f"Arrival"
                doc.add_heading(
                    f"Arrival: {date_string} {weekday}", 1)
                self.meals = "D/O"
            elif day_id == len(days) - 1 and dest_id == len(self.destinations) - 1:
                self.type = "Departure"
                doc.add_heading(f"Departure: {date_string} {weekday}", 1)
                self.meals = "B/L"
            else:
                self.type = "Tour"
                doc.add_heading(
                    f"Day {day_name} {date_string} {weekday}", 1)
        except Exception as e:
            print(e)
            raise e

    def gen_experiences(self, doc, dest_name, experiences):
        try:
            doc.add_heading(f"Experiences", 2)
            for exp_index, exp_name in enumerate(experiences):
                experience = experiences[exp_name]
                heading = doc.add_heading(exp_name, 3)
                run = heading.add_run()
                value = experience.get("value")
                description = value.get("description")
                for key, value in value.items():
                    if self.check_space(str(key)) >= 1 and value is None:
                        description += " "+key
                paragraph = doc.add_paragraph(
                    description, style='List Paragraph'
                )
                run = self.set_font(paragraph, doc)
                image = os.path.join(
                    self.file_dir, (f"images\\{dest_name}.png"))
                doc.add_picture(image, width=Inches(4),
                                height=Inches(4))
                self.gen_next(exp_index, doc, experiences)
        except Exception as e:
            print(e)
            raise e

    def check_space(self, string):
        count = 0
        for i in range(0, len(string)):
            if string[i] == " ":
                count += 1
        return count

    def gen_next(self, exp_index, doc, experiences):
        try:
            if exp_index < len(experiences) - 1:
                next = list(experiences)[exp_index + 1]
                next_description = experiences[next]['description'].split('.')[
                    0]
                paragraph = doc.add_paragraph(
                    f"Next we're going to get, {next_description}", style='List Paragraph'
                )
                run = self.set_font(paragraph, doc)
            else:
                if self.type == "Departure":
                    paragraph = doc.add_paragraph(
                        f"Take any photos to go back to Home", style='List Paragraph'
                    )
                    run = self.set_font(paragraph, doc)
                else:
                    paragraph = doc.add_paragraph(
                        f"Next we're going back to hotel to rest and take a meal", style='List Paragraph'
                    )
                    run = self.set_font(paragraph, doc)
                paragraph = doc.add_paragraph(
                    f"{self.meals}", style='List Paragraph')
                run = self.set_font(paragraph, doc)
        except Exception as e:
            print(e)
            raise e

    def set_font(self, text, doc, style="List Paragraph", family="Times New Roman", size=10):
        try:
            style = doc.styles[style]
            font = style.font
            font.name = family
            font.size = Pt(size)
            text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            text.style = style
            run = text.add_run()
            return run
        except Exception as e:
            print(e)
            raise e

    def gen_doc(self):
        try:
            new_document = Document()
            composer = Composer(new_document)
            self.docs = list(self.docs)
            self.docs.sort()
            new_list = [None] * len(self.docs)
            for doc_index, document in enumerate(self.docs):
                if "cover" in document:
                    new_list[0] = document
                else:
                    id = int(document.split("-")[1:][6].split(".")[0])+1
                    new_list[id] = document
            for doc_index, document in enumerate(new_list):
                doc = Document(document)
                if doc_index != len(new_list) - 1:
                    doc.add_page_break()
                composer.append(doc)
            composer.save(self.composed)
            convert(self.composed, self.pdf)
            print(f"DOC {self.travel_code} SAVED")
        except Exception as e:
            print(e)
            raise e

    def send_email_doc(self):
        try:
            self.net.sender.connect()
            self.net.sender.sendmail(
                self.email, f"Quote {self.travel_code}", "", [self.composed, self.pdf])
            print(f"DOC {self.travel_code} SENT")
        except Exception as e:
            print(e)
            raise e

    def del_tmp_folder(self):
        try:
            for filename in os.listdir(self.tempdir):
                file_path = os.path.join(self.tempdir, filename)
                self.delete_file(file_path)
        except Exception as e:
            print(e)
            raise e

    def delete_file(self, file_path):
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print('Failed to delete %s. Reason: %s' % (file_path, e))
            raise e

    def replace_word(self, document, key, word):
        try:
            for paragraph in document.paragraphs:
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, word)
        except Exception as e:
            print(e)
            raise e

    def check_date(self, date, format):
        try:
            res = datetime.strptime(date, format)
            return res
        except Exception as e:
            print(e)
            return date

    def delete_doc(self, file):
        try:
            if os.path.exists(file):
                os.remove(file)
        except Exception as e:
            print(e)


def test():
    try:
        file_dir = os.path.dirname(__file__)
        data_path = os.path.join(file_dir, ("data.json"))
        test_data = open(data_path, "r", encoding="utf8")
        data = json.load(test_data)
        data = dict(yaml.safe_load(data["data"]))
        doc_generator = DocGenerator(data)
        doc_generator.gen_tour_doc()
    except Exception as e:
        print(e)


if __name__ == "__main__":
    test()
